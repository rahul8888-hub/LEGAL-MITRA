import os
import json
import google.generativeai as genai
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyPDF2 import PdfReader
from io import BytesIO
import base64
from datetime import datetime

# Set page configuration (must be the first Streamlit command)
st.set_page_config(
    page_title="Indian Legal Document Generator",
    page_icon="⚖️",
    layout="wide"
)

# Configure Streamlit theme to maintain consistency in deployment
st.markdown("""
    <style>
        /* Global theme settings to override Streamlit defaults */
        :root {
            --background-color: white;
            --text-color: black;
            --header-color: #4F46E5;
        }
        
        /* Override Streamlit elements with !important flags for deployment */
        .stApp {
            background-color: white !important;
        }
        
        .main {
            background-color: white !important;
            color: black !important;
        }
        
        h1, h2, h3, h4, h5, h6 {
            color: #4F46E5 !important;
        }
        
        /* Force all input fields to maintain white background and black text */
        .stTextInput > div > div > input, 
        .stTextArea > div > div > textarea,
        .stNumberInput > div > div > input,
        .stDateInput > div > div > input,
        .stSelectbox > div,
        .stMultiselect > div {
            background-color: white !important;
            color: black !important;
            border: 1px solid #ccc !important;
        }
        
        /* Strong override for select dropdowns */
        .stSelectbox > div > div,
        .stSelectbox > div > div > div,
        .stSelectbox select {
            background-color: white !important;
            color: black !important;
        }
        
        /* Ensure all labels are black and visible */
        .stTextInput label, 
        .stTextArea label, 
        .stNumberInput label, 
        .stDateInput label, 
        .stSelectbox label,
        .stMultiselect label {
            color: black !important;
            font-weight: 500 !important;
        }
        
        /* Button styling - maintain consistent look */
        .stButton > button {
            background-color: #4F46E5 !important;
            color: white !important;
            border-radius: 5px !important;
            border: none !important;
            padding: 0.5rem 1rem !important;
        }
        
        .stButton > button:hover {
            background-color: #4338CA !important;
        }
        
        /* Markdown text color */
        .stMarkdown {
            color: black !important;
        }
        
        /* Placeholder text */
        ::placeholder {
            color: #666 !important;
            opacity: 1 !important;
        }
        
        /* Force override any Streamlit Cloud theme settings */
        [data-testid="stForm"] {
            background-color: white !important;
            border-color: #e0e0e0 !important;
        }
        
        /* Text inside widgets - additional selector for deployment */
        [data-baseweb="select"] > div,
        [data-baseweb="input"] > div {
            background-color: white !important;
            color: black !important;
        }
    </style>
""", unsafe_allow_html=True)

# Configure the Gemini API
GEMINI_API_KEY = "AIzaSyCTD2PAIDkxQcx4dy8BUBTr2ER6QVVNjAo"
genai.configure(api_key=GEMINI_API_KEY)

# Path to templates folder
TEMPLATES_FOLDER = "data"

# Dictionary of template information
TEMPLATES = {
    "divorce_petition": {
        "title": "Divorce Petition",
        "description": "Legal application for divorce filing",
        "filename": "Divorce Petition.txt",
        "placeholders": [
            "court_name", "petition_number", "year",
            "petitioner_name", "petitioner_father_name", "petitioner_address",
            "respondent_name", "respondent_father_name", "respondent_address",
            "marriage_date", "marriage_place", "separation_date",
            "child_name", "child_dob", "grounds_for_divorce",
            "dowry_amount1", "dowry_amount2", "abuse_details",
            "financial_demand", "panchayat_date1", "panchayat_date2",
            "witness1_name", "witness1_father", "witness1_address",
            "witness2_name", "witness2_father", "witness2_address",
            "jurisdiction_details", "court_fee_details",
            "verification_place", "verification_date", "advocate_name"
        ]
    },
    "rental_lease_agreement": {
        "title": "Rental Lease Agreement (Residential)",
        "description": "Contract between landlord and tenant for residential property",
        "filename": "Rental Lease Agreement (Residential).txt",
        "placeholders": [
            "landlord_name", "landlord_address", "landlord_pan",
            "tenant_name", "tenant_address", "tenant_pan",
            "property_address", "property_type", "property_area",
            "lease_start_date", "lease_end_date", "lease_duration",
            "monthly_rent", "rent_due_date", "payment_method",
            "security_deposit", "maintenance_terms", "utilities_responsibility",
            "termination_notice_period", "termination_conditions",
            "witness1_name", "witness1_address", "witness2_name", "witness2_address",
            "execution_date", "execution_place", "stamp_duty_details"
        ]
    },
    "power_of_attorney": {
        "title": "Power of Attorney",
        "description": "Legal authorization to act on another's behalf",
        "filename": "Power of Attorney.txt",
        "placeholders": [
            "principal_name", "principal_father_name", "principal_address",
            "principal_pan", "principal_aadhar", "principal_occupation",
            "attorney_name", "attorney_father_name", "attorney_address",
            "attorney_pan", "attorney_aadhar", "attorney_occupation",
            "powers_granted", "property_details", "bank_accounts_details",
            "effective_date", "termination_conditions",
            "witness1_name", "witness1_address", "witness2_name", "witness2_address",
            "execution_date", "execution_place", "stamp_duty_details"
        ]
    },
    "will_testament": {
        "title": "Will and Testament",
        "description": "Legal document expressing how one wishes their property to be distributed",
        "filename": "Will and Testament.txt",
        "placeholders": [
            "testator_name", "testator_father_name", "testator_address",
            "testator_age", "testator_occupation", "testator_religion",
            "spouse_name", "children_names", "children_addresses",
            "property_details", "bank_accounts", "investments",
            "specific_bequests", "residuary_clause",
            "executor_name", "executor_address", "alternate_executor_name",
            "witness1_name", "witness1_address", "witness2_name", "witness2_address",
            "execution_date", "execution_place", "stamp_duty_details"
        ]
    },
    "partnership_deed": {
        "title": "Partnership Deed",
        "description": "Legal agreement between business partners",
        "filename": "Partnership Deed.txt",
        "placeholders": [
            "firm_name", "business_address", "business_nature",
            "partner1_name", "partner1_father_name", "partner1_address", "partner1_age",
            "partner2_name", "partner2_father_name", "partner2_address", "partner2_age",
            "partner3_name", "partner3_father_name", "partner3_address", "partner3_age",
            "partner4_name", "partner4_father_name", "partner4_address", "partner4_age",
            "capital_contribution", "profit_sharing_ratio", "interest_rate",
            "working_partners", "remuneration_details", "financial_year_end",
            "bank_operations", "dispute_resolution", "arbitration_details",
            "witness1_name", "witness1_address", "witness2_name", "witness2_address",
            "execution_date", "execution_place", "stamp_duty_details"
        ]
    },
    "mou": {
        "title": "Memorandum of Understanding (MOU)",
        "description": "Formal agreement between two or more parties",
        "filename": "Memorandum of Understanding (MOU).txt",
        "placeholders": [
            "party1_name", "party1_address", "party1_representative",
            "party2_name", "party2_address", "party2_representative",
            "purpose", "scope", "responsibilities",
            "timeframe", "confidentiality_terms", "termination_clause",
            "dispute_resolution", "governing_law",
            "witness1_name", "witness1_address", "witness2_name", "witness2_address",
            "execution_date", "execution_place", "stamp_duty_details"
        ]
    },
    "sale_deed": {
        "title": "Sale Deed",
        "description": "Legal document for transferring property ownership",
        "filename": "Sale Deed.txt",
        "placeholders": [
            "seller_name", "seller_father_name", "seller_address",
            "seller_pan", "seller_aadhar", "seller_occupation",
            "buyer_name", "buyer_father_name", "buyer_address",
            "buyer_pan", "buyer_aadhar", "buyer_occupation",
            "property_description", "property_address", "property_area",
            "sale_amount", "payment_terms", "possession_date",
            "encumbrance_details", "title_clearance",
            "witness1_name", "witness1_address", "witness2_name", "witness2_address",
            "execution_date", "execution_place", "stamp_duty_details"
        ]
    },
    "rent_agreement_commercial": {
        "title": "Rent Agreement (Commercial)",
        "description": "Contract for commercial property rental",
        "filename": "Rent Agreement(Commerical).txt",
        "placeholders": [
            "landlord_name", "landlord_address", "landlord_pan",
            "tenant_name", "tenant_address", "tenant_pan",
            "property_address", "property_type", "property_area",
            "rent_amount", "lease_period", "security_deposit",
            "usage_purpose", "maintenance_terms", "renewal_terms",
            "termination_conditions", "governing_law",
            "witness1_name", "witness1_address", "witness2_name", "witness2_address",
            "execution_date", "execution_place", "stamp_duty_details"
        ]
    }
}

def read_template_file(template_type):
    template = TEMPLATES[template_type]
    file_path = os.path.join(TEMPLATES_FOLDER, template["filename"])
    
    if not os.path.exists(file_path):
        return None
    
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.txt':
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    elif file_extension == '.docx':
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    elif file_extension == '.pdf':
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    else:
        return None

def generate_document_content(template_type, user_inputs):
    template_content = read_template_file(template_type)
    
    if not template_content:
        return "Error: Template file not found or unsupported format."
    
    user_inputs["current_date"] = datetime.now().strftime("%d %B, %Y")
    
    prompt = f"""
    I have an Indian legal document template for {TEMPLATES[template_type]['title']} with the following content:
    
    {template_content}
    
    Please generate a complete and legally appropriate version of this document using the following information:
    
    {json.dumps(user_inputs, indent=2)}
    
    Fill in all the necessary details based on the provided information.
    Maintain the proper structure and legal language of Indian legal documents.
    Include all required clauses and sections.
    Format the document professionally.
    Return ONLY the completed document with no additional commentary.
    """
    
    model = genai.GenerativeModel('gemini-1.5-pro')
    response = model.generate_content(prompt)
    
    return response.text

def create_word_document(content, template_type):
    doc = Document()
    
    title = doc.add_heading(TEMPLATES[template_type]["title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    paragraphs = content.split('\n')
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph(para)
            for run in p.runs:
                run.font.size = Pt(12)
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def get_download_link(doc_io, filename):
    b64 = base64.b64encode(doc_io.getvalue()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Document</a>'

def main():
    st.title("Indian Legal Document Generator")
    st.markdown("Generate customized Indian legal documents based on your inputs")
    
    template_type = st.selectbox(
        "Select Document Type",
        options=list(TEMPLATES.keys()),
        format_func=lambda x: TEMPLATES[x]["title"]
    )
    
    st.write(f"Description: {TEMPLATES[template_type]['description']}")
    
    template_file = os.path.join(TEMPLATES_FOLDER, TEMPLATES[template_type]["filename"])
    if not os.path.exists(template_file):
        st.warning(f"Template file not found: {TEMPLATES[template_type]['filename']}. Please make sure the file exists in the '{TEMPLATES_FOLDER}' folder.")
    
    with st.form(key=f"template_form_{template_type}"):
        user_inputs = {}
        
        placeholders = TEMPLATES[template_type]["placeholders"]
        
        col1, col2 = st.columns(2)
        
        for i, placeholder in enumerate(placeholders):
            label = " ".join(word.capitalize() for word in placeholder.split("_"))
            
            with col1 if i % 2 == 0 else col2:
                if any(term in placeholder for term in ["description", "details", "terms", "clause", "statements", "conditions"]):
                    user_inputs[placeholder] = st.text_area(label)
                elif "date" in placeholder:
                    user_inputs[placeholder] = st.date_input(label).strftime("%d/%m/%Y")
                elif any(term in placeholder for term in ["amount", "deposit", "rate"]):
                    user_inputs[placeholder] = st.number_input(label, min_value=0, step=1000)
                else:
                    user_inputs[placeholder] = st.text_input(label)
        
        submit_button = st.form_submit_button(label="Generate Document")
    
    if submit_button:
        if all(user_inputs.values()):
            with st.spinner("Generating document..."):
                try:
                    document_content = generate_document_content(template_type, user_inputs)
                    
                    doc_io = create_word_document(document_content, template_type)
                    
                    st.subheader("Generated Document")
                    st.write(document_content)
                    
                    filename = f"{TEMPLATES[template_type]['title'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    st.markdown(get_download_link(doc_io, filename), unsafe_allow_html=True)
                    
                except Exception as e:
                    st.error(f"Error generating document: {str(e)}")
        else:
            st.warning("Please fill in all fields")

if __name__ == "__main__":
    main()