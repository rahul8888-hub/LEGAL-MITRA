import os
import json
from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyPDF2 import PdfReader
from io import BytesIO
import base64
from datetime import datetime
import google.generativeai as genai

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

# Configure the Gemini API
GEMINI_API_KEY = "AIzaSyCTD2PAIDkxQcx4dy8BUBTr2ER6QVVNjAo"
genai.configure(api_key=GEMINI_API_KEY)

# Path to templates folder
TEMPLATES_FOLDER = "data"

# Dictionary of template information
TEMPLATES = {
    "divorce_petition": {
        "title": "Divorce Petition",
        "description": "Legal application for divorce filing",
        "filename": "Divorce Petition.txt",
        "placeholders": [
            "court_name", "petition_number", "year",
            "petitioner_name", "petitioner_father_name", "petitioner_address",
            "respondent_name", "respondent_father_name", "respondent_address",
            "marriage_date", "marriage_place", "separation_date",
            "child_name", "child_dob", "grounds_for_divorce",
            "dowry_amount1", "dowry_amount2", "abuse_details",
            "financial_demand", "panchayat_date1", "panchayat_date2",
            "witness1_name", "witness1_father", "witness1_address",
            "witness2_name", "witness2_father", "witness2_address",
            "jurisdiction_details", "court_fee_details",
            "verification_place", "verification_date", "advocate_name"
        ]
    },
    "rental_lease_agreement": {
        "title": "Rental Lease Agreement (Residential)",
        "description": "Contract between landlord and tenant for residential property",
        "filename": "Rental Lease Agreement (Residential).txt",
        "placeholders": [
            "landlord_name", "landlord_address", "landlord_pan",
            "tenant_name", "tenant_address", "tenant_pan",
            "property_address", "property_type", "property_area",
            "lease_start_date", "lease_end_date", "lease_duration",
            "monthly_rent", "rent_due_date", "payment_method",
            "security_deposit", "maintenance_terms", "utilities_responsibility",
            "termination_notice_period", "termination_conditions",
            "witness1_name", "witness1_address", "witness2_name", "witness2_address",
            "execution_date", "execution_place", "stamp_duty_details"
        ]
    },
    "power_of_attorney": {
        "title": "Power of Attorney",
        "description": "Legal authorization to act on another's behalf",
        "filename": "Power of Attorney.txt",
        "placeholders": [
            "principal_name", "principal_father_name", "principal_address",
            "principal_pan", "principal_aadhar", "principal_occupation",
            "attorney_name", "attorney_father_name", "attorney_address",
            "attorney_pan", "attorney_aadhar", "attorney_occupation",
            "powers_granted", "property_details", "bank_accounts_details",
            "effective_date", "termination_conditions",
            "witness1_name", "witness1_address", "witness2_name", "witness2_address",
            "execution_date", "execution_place", "stamp_duty_details"
        ]
    },
    "will_testament": {
        "title": "Will and Testament",
        "description": "Legal document expressing how one wishes their property to be distributed",
        "filename": "Will and Testament.txt",
        "placeholders": [
            "testator_name", "testator_father_name", "testator_address",
            "testator_age", "testator_occupation", "testator_religion",
            "spouse_name", "children_names", "children_addresses",
            "property_details", "bank_accounts", "investments",
            "specific_bequests", "residuary_clause",
            "executor_name", "executor_address", "alternate_executor_name",
            "witness1_name", "witness1_address", "witness2_name", "witness2_address",
            "execution_date", "execution_place", "stamp_duty_details"
        ]
    },
    "partnership_deed": {
        "title": "Partnership Deed",
        "description": "Legal agreement between business partners",
        "filename": "Partnership Deed.txt",
        "placeholders": [
            "firm_name", "business_address", "business_nature",
            "partner1_name", "partner1_father_name", "partner1_address", "partner1_age",
            "partner2_name", "partner2_father_name", "partner2_address", "partner2_age",
            "partner3_name", "partner3_father_name", "partner3_address", "partner3_age",
            "partner4_name", "partner4_father_name", "partner4_address", "partner4_age",
            "capital_contribution", "profit_sharing_ratio", "interest_rate",
            "working_partners", "remuneration_details", "financial_year_end",
            "bank_operations", "dispute_resolution", "arbitration_details",
            "witness1_name", "witness1_address", "witness2_name", "witness2_address",
            "execution_date", "execution_place", "stamp_duty_details"
        ]
    },
    "mou": {
        "title": "Memorandum of Understanding (MOU)",
        "description": "Formal agreement between two or more parties",
        "filename": "Memorandum of Understanding (MOU).txt",
        "placeholders": [
            "party1_name", "party1_address", "party1_representative",
            "party2_name", "party2_address", "party2_representative",
            "purpose", "scope", "responsibilities",
            "timeframe", "confidentiality_terms", "termination_clause",
            "dispute_resolution", "governing_law",
            "witness1_name", "witness1_address", "witness2_name", "witness2_address",
            "execution_date", "execution_place", "stamp_duty_details"
        ]
    },
    "sale_deed": {
        "title": "Sale Deed",
        "description": "Legal document for transferring property ownership",
        "filename": "Sale Deed.txt",
        "placeholders": [
            "seller_name", "seller_father_name", "seller_address",
            "seller_pan", "seller_aadhar", "seller_occupation",
            "buyer_name", "buyer_father_name", "buyer_address",
            "buyer_pan", "buyer_aadhar", "buyer_occupation",
            "property_description", "property_address", "property_area",
            "sale_amount", "payment_terms", "possession_date",
            "encumbrance_details", "title_clearance",
            "witness1_name", "witness1_address", "witness2_name", "witness2_address",
            "execution_date", "execution_place", "stamp_duty_details"
        ]
    },
    "rent_agreement_commercial": {
        "title": "Rent Agreement (Commercial)",
        "description": "Contract for commercial property rental",
        "filename": "Rent Agreement(Commerical).txt",
        "placeholders": [
            "landlord_name", "landlord_address", "landlord_pan",
            "tenant_name", "tenant_address", "tenant_pan",
            "property_address", "property_type", "property_area",
            "rent_amount", "lease_period", "security_deposit",
            "usage_purpose", "maintenance_terms", "renewal_terms",
            "termination_conditions", "governing_law",
            "witness1_name", "witness1_address", "witness2_name", "witness2_address",
            "execution_date", "execution_place", "stamp_duty_details"
        ]
    }
}

def read_template_file(template_type):
    template = TEMPLATES[template_type]
    file_path = os.path.join(TEMPLATES_FOLDER, template["filename"])
    
    if not os.path.exists(file_path):
        return None
    
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.txt':
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    elif file_extension == '.docx':
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    elif file_extension == '.pdf':
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    else:
        return None

def generate_document_content(template_type, user_inputs):
    template_content = read_template_file(template_type)
    
    if not template_content:
        return "Error: Template file not found or unsupported format."
    
    user_inputs["current_date"] = datetime.now().strftime("%d %B, %Y")
    
    prompt = f"""
    I have an Indian legal document template for {TEMPLATES[template_type]['title']} with the following content:
    
    {template_content}
    
    Please generate a complete and legally appropriate version of this document using the following information:
    
    {json.dumps(user_inputs, indent=2)}
    
    Fill in all the necessary details based on the provided information.
    Maintain the proper structure and legal language of Indian legal documents.
    Include all required clauses and sections.
    Format the document professionally.
    Return ONLY the completed document with no additional commentary.
    """
    
    model = genai.GenerativeModel('gemini-1.5-pro')
    response = model.generate_content(prompt)
    
    return response.text

def create_word_document(content, template_type):
    doc = Document()
    
    title = doc.add_heading(TEMPLATES[template_type]["title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    paragraphs = content.split('\n')
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph(para)
            for run in p.runs:
                run.font.size = Pt(12)
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    # Convert to base64 for download
    encoded_document = base64.b64encode(doc_io.getvalue()).decode('utf-8')
    
    return encoded_document

@app.route('/api/templates', methods=['GET'])
def get_templates():
    templates_info = {}
    for key, template in TEMPLATES.items():
        templates_info[key] = {
            "title": template["title"],
            "description": template["description"],
            "placeholders": template["placeholders"]
        }
    
    return jsonify(templates_info)

@app.route('/api/generate-document', methods=['POST'])
def generate_document():
    data = request.json
    template_type = data.get('template_type')
    user_inputs = data.get('user_inputs')
    
    if not template_type or not user_inputs:
        return jsonify({"error": "Missing template_type or user_inputs"}), 400
    
    if template_type not in TEMPLATES:
        return jsonify({"error": f"Invalid template type: {template_type}"}), 400
    
    try:
        document_content = generate_document_content(template_type, user_inputs)
        encoded_document = create_word_document(document_content, template_type)
        
        filename = f"{TEMPLATES[template_type]['title'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        return jsonify({
            "content": document_content,
            "document": encoded_document,
            "filename": filename
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5005, debug=True)